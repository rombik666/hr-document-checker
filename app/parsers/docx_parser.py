from pathlib import Path
from uuid import uuid4
from datetime import datetime, timezone

from docx import Document

from app.parsers.base import DocumentParser
from app.schemas.common import DocumentType, ProcessingStatus, SourceFormat, StorageMode
from app.schemas.documents import DocumentMetadata, DocumentSection, ParsedDocument


class DOCXParser(DocumentParser):
    """
    Парсер DOCX-документов.

    На этом этапе он делает три вещи:
    1. Извлекает обычный текст из абзацев.
    2. Извлекает текст из таблиц.
    3. Формирует список секций документа.

    Позже поверх этих секций мы добавим:
    - определение типа секции: contacts, experience, skills, education;
    - извлечение контактов;
    - извлечение дат;
    - классификацию типа документа.
    """

    def parse(self, file_path: Path) -> ParsedDocument:
        if not file_path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        if file_path.suffix.lower() != ".docx":
            raise ValueError(f"DOCXParser поддерживает только .docx, получено: {file_path.suffix}")

        document = Document(file_path)

        sections: list[DocumentSection] = []
        text_parts: list[str] = []

        position = 0

        # 1. Читаем обычные абзацы
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if not text:
                continue

            section = DocumentSection(
                section_id=str(uuid4()),
                section_type="paragraph",
                title=None,
                text=text,
                position_in_document=position,
                metadata={
                    "source": "paragraph",
                    "style": paragraph.style.name if paragraph.style else None,
                },
            )

            sections.append(section)
            text_parts.append(text)
            position += 1

        # 2. Читаем таблицы
        for table_index, table in enumerate(document.tables):
            rows_as_text: list[str] = []

            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]

                if cells:
                    rows_as_text.append(" | ".join(cells))

            if not rows_as_text:
                continue

            table_text = "\n".join(rows_as_text)

            section = DocumentSection(
                section_id=str(uuid4()),
                section_type="table",
                title=f"Таблица {table_index + 1}",
                text=table_text,
                position_in_document=position,
                metadata={
                    "source": "table",
                    "table_index": table_index,
                },
            )

            sections.append(section)
            text_parts.append(table_text)
            position += 1

        raw_text = "\n\n".join(text_parts)

        metadata = DocumentMetadata(
            document_id=str(uuid4()),
            document_type=DocumentType.UNKNOWN,
            source_format=SourceFormat.DOCX,
            filename=file_path.name,
            upload_time=datetime.now(timezone.utc),
            processing_status=ProcessingStatus.PARSED,
            storage_mode=StorageMode.TEMPORARY,
            hash=None,
            warnings=[],
        )

        return ParsedDocument(
            metadata=metadata,
            raw_text=raw_text,
            sections=sections,
            entities=[],
        )