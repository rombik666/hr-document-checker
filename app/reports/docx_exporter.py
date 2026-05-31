from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.checks import Issue
from app.schemas.common import ReportStatus, Severity
from app.schemas.reports import Report


class DocxReportExporter:

    def export(
        self,
        report: Report,
        include_technical: bool = True,
    ) -> BytesIO:
        document = Document()

        self._configure_styles(document)

        self._add_title(document, report, include_technical=include_technical)
        self._add_summary(document, report)
        self._add_counts(document, report)
        self._add_vacancy_relevance(document, report)

        if include_technical:
            self._add_issues_section(document, "Critical", report.critical, include_technical=True)
            self._add_issues_section(document, "Major", report.major, include_technical=True)
            self._add_issues_section(document, "Minor", report.minor, include_technical=True)
            self._add_technical_info(document, report)
        else:
            self._add_issues_section(document, "Критические замечания", report.critical, include_technical=False)
            self._add_issues_section(document, "Значимые замечания", report.major, include_technical=False)
            self._add_issues_section(document, "Незначительные замечания", report.minor, include_technical=False)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        return buffer

    @staticmethod
    def _configure_styles(document: Document) -> None:
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(12)

        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = document.styles[style_name]
            style.font.name = "Times New Roman"

    @staticmethod
    def _add_title(
        document: Document,
        report: Report,
        include_technical: bool,
    ) -> None:
        title = document.add_heading("Итоговый отчёт проверки документа", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph(f"Файл: {report.filename}")

        if include_technical:
            document.add_paragraph(f"Report ID: {report.report_id}")
            document.add_paragraph(f"Document ID: {report.document_id}")

    @staticmethod
    def _status_label(status: ReportStatus | str) -> str:
        value = getattr(status, "value", status)

        return {
            "ready": "Готов",
            "requires_revision": "Требует доработки",
            "failed": "Ошибка проверки",
            "draft": "Черновик",
        }.get(str(value), str(value))

    @staticmethod
    def _add_summary(document: Document, report: Report) -> None:
        document.add_heading("1. Общий результат", level=2)

        document.add_paragraph(
            f"Статус: {DocxReportExporter._status_label(report.summary_status)}"
        )
        document.add_paragraph(report.summary)

    @staticmethod
    def _add_counts(document: Document, report: Report) -> None:
        document.add_heading("2. Сводка по замечаниям", level=2)

        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        header_cells[0].text = "Всего"
        header_cells[1].text = "Критические"
        header_cells[2].text = "Значимые"
        header_cells[3].text = "Незначительные"

        row_cells = table.add_row().cells
        row_cells[0].text = str(report.total_issues)
        row_cells[1].text = str(report.critical_count)
        row_cells[2].text = str(report.major_count)
        row_cells[3].text = str(report.minor_count)

    @staticmethod
    def _add_vacancy_relevance(document: Document, report: Report) -> None:
        if report.vacancy_relevance is None:
            return

        document.add_heading("3. Релевантность вакансии", level=2)

        relevance = report.vacancy_relevance

        document.add_paragraph(
            f"Покрытие требований: {relevance.coverage_percent}%"
        )

        if relevance.comment:
            document.add_paragraph(f"Комментарий: {relevance.comment}")

        document.add_paragraph("Закрытые требования:")

        if relevance.covered_requirements:
            for item in relevance.covered_requirements:
                document.add_paragraph(item, style="List Bullet")
        else:
            document.add_paragraph("Нет данных.")

        document.add_paragraph("Незакрытые требования:")

        if relevance.missing_requirements:
            for item in relevance.missing_requirements:
                document.add_paragraph(item, style="List Bullet")
        else:
            document.add_paragraph("Явные пробелы не найдены.")

    def _add_issues_section(
        self,
        document: Document,
        title: str,
        issues: list[Issue],
        include_technical: bool,
    ) -> None:
        document.add_heading(title, level=2)

        if not issues:
            document.add_paragraph("Замечаний не найдено.")
            return

        for index, issue in enumerate(issues, start=1):
            self._add_issue(
                document=document,
                index=index,
                issue=issue,
                include_technical=include_technical,
            )

    @staticmethod
    def _severity_label(severity: Severity | str) -> str:
        value = getattr(severity, "value", severity)

        return {
            "Critical": "Критическое",
            "Major": "Значимое",
            "Minor": "Незначительное",
            "critical": "Критическое",
            "major": "Значимое",
            "minor": "Незначительное",
        }.get(str(value), str(value))

    @staticmethod
    def _add_issue(
        document: Document,
        index: int,
        issue: Issue,
        include_technical: bool,
    ) -> None:
        if include_technical:
            title = f"{index}. {issue.issue_type} — {issue.severity.value}"
        else:
            title = f"{index}. {issue.issue_type} — {issue.description}"

        document.add_heading(title, level=3)

        if include_technical:
            document.add_paragraph(f"Агент: {issue.source_agent}")
            document.add_paragraph(f"Описание: {issue.description}")

        if issue.evidence_fragment:
            document.add_paragraph("Фрагмент:")
            document.add_paragraph(issue.evidence_fragment)

        if issue.recommendation:
            document.add_paragraph(
                f"Рекомендация: {issue.recommendation.recommendation_text}"
            )

            if issue.recommendation.example_fix:
                document.add_paragraph(
                    f"Пример исправления: {issue.recommendation.example_fix}"
                )

        if include_technical and issue.confidence_score is not None:
            document.add_paragraph(f"Confidence score: {issue.confidence_score}")

    @staticmethod
    def _add_technical_info(document: Document, report: Report) -> None:
        document.add_heading("Техническая справка", level=2)

        technical = report.technical_info

        document.add_paragraph(f"Дата генерации: {technical.generated_at}")
        document.add_paragraph(f"Количество агентов: {technical.total_agents_count}")
        document.add_paragraph(
            f"Успешных агентов: {technical.successful_agents_count}"
        )
        document.add_paragraph(f"Ошибок агентов: {technical.failed_agents_count}")

        document.add_paragraph("Выполненные проверки:")

        for check_name in technical.checks_completed:
            document.add_paragraph(check_name, style="List Bullet")

        document.add_paragraph("Метаданные:")

        for key, value in technical.metadata.items():
            document.add_paragraph(f"{key}: {value}")