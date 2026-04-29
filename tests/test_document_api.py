from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from tests.auth_helpers import auth_headers

from app.main import app


client = TestClient(app)


def test_parse_docx_endpoint_returns_parsed_document(tmp_path: Path) -> None:

    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Иван Иванов")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Навыки: Python, FastAPI, PostgreSQL")
    document.save(file_path)

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/parse",
            headers=auth_headers(client, "candidate"),
            files={
                "file": (
                    "resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200

    data = response.json()

    assert data["metadata"]["filename"] == "resume.docx"
    assert data["metadata"]["source_format"] == "docx"
    assert data["metadata"]["processing_status"] == "parsed"

    assert "Иван Иванов" in data["raw_text"]
    assert "ivan@example.com" in data["raw_text"]
    assert "Python" in data["raw_text"]

    assert len(data["sections"]) == 3

    # Секции уже прошли классификацию.
    # Исходный тип paragraph сохраняется в metadata.
    assert data["sections"][0]["metadata"]["original_section_type"] == "paragraph"

    entity_types = {entity["entity_type"] for entity in data["entities"]}
    skills = {
        entity["normalized_value"]
        for entity in data["entities"]
        if entity["entity_type"] == "skill"
    }

    assert "email" in entity_types
    assert "skill" in entity_types

    assert "python" in skills
    assert "fastapi" in skills
    assert "postgresql" in skills


def test_parse_endpoint_rejects_unsupported_file_format(tmp_path: Path) -> None:
    """
    Проверяем, что endpoint отклоняет неподдерживаемые форматы.
    """

    file_path = tmp_path / "resume.txt"
    file_path.write_text("Simple text file", encoding="utf-8")

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/parse",
            headers=auth_headers(client, "candidate"),
            files={
                "file": (
                    "resume.txt",
                    file,
                    "text/plain",
                )
            },
        )

    assert response.status_code == 400

    data = response.json()

    assert "Поддерживаются только файлы .docx и .pdf" in data["detail"]