from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


client = TestClient(app)


def test_saved_report_can_be_exported_to_docx(tmp_path: Path) -> None:
    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, Git")
    document.add_paragraph("Опыт работы:")
    document.add_paragraph("Backend developer, 2023-2024")
    document.add_paragraph("Занимался разработкой backend.")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/report",
            data={
                "vacancy_text": "Требования: Python, FastAPI, PostgreSQL, Docker, Git.",
                "storage_mode": "temporary",
            },
            files={
                "file": (
                    "resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200

    report = response.json()
    report_id = report["report_id"]

    export_response = client.get(
        f"/api/v1/documents/reports/{report_id}/export/docx"
    )

    assert export_response.status_code == 200

    content_type = export_response.headers["content-type"]

    assert content_type.startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    assert "attachment" in export_response.headers["content-disposition"]
    assert f"report_{report_id}.docx" in export_response.headers["content-disposition"]

    exported_document = Document(BytesIO(export_response.content))

    exported_text = "\n".join(
        paragraph.text
        for paragraph in exported_document.paragraphs
    )

    assert "Итоговый отчёт проверки документа" in exported_text
    assert "resume.docx" in exported_text
    assert "weak_phrase" in exported_text
    assert "vacancy_requirements_gap" in exported_text


def test_no_store_report_cannot_be_exported_to_docx(tmp_path: Path) -> None:
    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, Git")
    document.add_paragraph("Опыт работы:")
    document.add_paragraph("Backend developer, 2023-2024")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/report",
            data={
                "storage_mode": "no_store",
            },
            files={
                "file": (
                    "resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200

    report = response.json()
    report_id = report["report_id"]

    export_response = client.get(
        f"/api/v1/documents/reports/{report_id}/export/docx"
    )

    assert export_response.status_code == 404