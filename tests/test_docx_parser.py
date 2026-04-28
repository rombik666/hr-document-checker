from pathlib import Path

from docx import Document

from app.parsers.docx_parser import DOCXParser
from app.schemas.common import ProcessingStatus, SourceFormat


def test_docx_parser_extracts_paragraphs(tmp_path: Path) -> None:

    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Иван Иванов")
    document.add_paragraph("Python-разработчик")
    document.add_paragraph("Опыт работы: 2 года")
    document.save(file_path)

    parser = DOCXParser()
    parsed_document = parser.parse(file_path)

    assert parsed_document.metadata.filename == "resume.docx"
    assert parsed_document.metadata.source_format == SourceFormat.DOCX
    assert parsed_document.metadata.processing_status == ProcessingStatus.PARSED

    assert "Иван Иванов" in parsed_document.raw_text
    assert "Python-разработчик" in parsed_document.raw_text
    assert "Опыт работы: 2 года" in parsed_document.raw_text

    assert len(parsed_document.sections) == 3
    assert parsed_document.sections[0].section_type == "paragraph"
    assert parsed_document.sections[0].text == "Иван Иванов"


def test_docx_parser_extracts_tables(tmp_path: Path) -> None:

    file_path = tmp_path / "resume_with_table.docx"

    document = Document()
    document.add_paragraph("Резюме")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Навык"
    table.cell(0, 1).text = "Уровень"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "Продвинутый"

    document.save(file_path)

    parser = DOCXParser()
    parsed_document = parser.parse(file_path)

    assert "Резюме" in parsed_document.raw_text
    assert "Навык | Уровень" in parsed_document.raw_text
    assert "Python | Продвинутый" in parsed_document.raw_text

    section_types = [section.section_type for section in parsed_document.sections]

    assert "paragraph" in section_types
    assert "table" in section_types


def test_docx_parser_raises_error_for_missing_file(tmp_path: Path) -> None:

    file_path = tmp_path / "missing.docx"

    parser = DOCXParser()

    try:
        parser.parse(file_path)
    except FileNotFoundError as error:
        assert "Файл не найден" in str(error)
    else:
        raise AssertionError("Ожидался FileNotFoundError")