from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


client = TestClient(app)


def test_metrics_endpoint_returns_snapshot() -> None:
    response = client.get("/api/v1/metrics")

    assert response.status_code == 200

    data = response.json()

    assert "requests_total" in data
    assert "documents_processed_total" in data
    assert "reports_generated_total" in data
    assert "errors_total" in data
    assert "issues_found_total" in data
    assert "average_request_time_ms" in data
    assert "average_document_processing_time_ms" in data


def test_report_generation_updates_metrics(tmp_path: Path) -> None:
    before_response = client.get("/api/v1/metrics")
    before = before_response.json()

    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, Git")
    document.add_paragraph("Опыт работы:")
    document.add_paragraph("Backend developer, 2023-2024")
    document.add_paragraph("Занимался разработкой backend.")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/report",
            data={
                "storage_mode": "no_store",
            },
            files={
                "file": (
                    "resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200

    after_response = client.get("/api/v1/metrics")
    after = after_response.json()

    assert after["documents_processed_total"] >= before["documents_processed_total"] + 1
    assert after["reports_generated_total"] >= before["reports_generated_total"] + 1
    assert after["requests_total"] >= before["requests_total"] + 2