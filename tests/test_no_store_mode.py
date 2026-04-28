from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


client = TestClient(app)


def test_report_endpoint_no_store_does_not_persist_report(tmp_path: Path) -> None:
    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, FastAPI, PostgreSQL")
    document.add_paragraph("Опыт работы: Python-разработчик, 2021-2024")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/report",
            data={
                "storage_mode": "no_store",
            },
            files={
                "file": (
                    "resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200

    report = response.json()

    assert report["technical_info"]["metadata"]["storage_mode"] == "no_store"
    assert report["technical_info"]["metadata"]["saved_to_db"] is False

    report_id = report["report_id"]

    get_response = client.get(f"/api/v1/documents/reports/{report_id}")

    assert get_response.status_code == 404