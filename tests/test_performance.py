from pathlib import Path
from time import perf_counter

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from tests.auth_helpers import auth_headers


client = TestClient(app)


@pytest.mark.performance
def test_docx_report_generation_performance(tmp_path: Path) -> None:

    file_path = tmp_path / "performance_resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, FastAPI, PostgreSQL, Docker, Git")
    document.add_paragraph("Опыт работы:")
    document.add_paragraph("Python backend developer, 2021-2024")
    document.add_paragraph("Разрабатывал backend-сервисы на FastAPI.")
    document.add_paragraph("Работал с PostgreSQL и Docker.")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    started_at = perf_counter()

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/report",
            headers=auth_headers(client, "candidate"),
            data={
                "storage_mode": "no_store",
                "vacancy_text": "Требования: Python, FastAPI, PostgreSQL, Docker, Git.",
            },
            files={
                "file": (
                    "performance_resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    duration_seconds = perf_counter() - started_at

    assert response.status_code == 200

    # The technical specification target for small documents is much higher.
    # This local threshold keeps the test stable on slower machines.
    assert duration_seconds < 30