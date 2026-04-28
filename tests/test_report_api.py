from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


client = TestClient(app)


def test_report_endpoint_returns_report_for_valid_docx(tmp_path: Path) -> None:
    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, FastAPI, PostgreSQL")
    document.add_paragraph("Опыт работы: Python-разработчик, 2021-2024")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/report",
            files={
                "file": (
                    "resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200

    data = response.json()

    assert data["filename"] == "resume.docx"
    assert data["summary_status"] == "ready"
    assert data["total_issues"] == 0

    assert data["critical"] == []
    assert data["major"] == []
    assert data["minor"] == []

    assert data["technical_info"]["total_agents_count"] == 7
    assert data["technical_info"]["metadata"]["document_type"] == "cv"
    assert "CompletenessAgent" in data["technical_info"]["checks_completed"]


def test_report_endpoint_rejects_unsupported_file_format(tmp_path: Path) -> None:
    file_path = tmp_path / "resume.txt"
    file_path.write_text("Simple text file", encoding="utf-8")

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/report",
            files={
                "file": (
                    "resume.txt",
                    file,
                    "text/plain",
                )
            },
        )

    assert response.status_code == 400

    data = response.json()

    assert "Поддерживаются только файлы .docx и .pdf" in data["detail"]


def test_report_endpoint_includes_semantic_issues_and_vacancy_relevance(tmp_path: Path) -> None:
    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, Git")
    document.add_paragraph("Опыт работы:")
    document.add_paragraph("Backend developer, 2023-2024")
    document.add_paragraph("Занимался разработкой backend.")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/report",
            data={
                "vacancy_text": "Требования: Python, FastAPI, PostgreSQL, Docker, Git.",
                "storage_mode": "temporary",
            },
            files={
                "file": (
                    "resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200

    data = response.json()

    assert data["filename"] == "resume.docx"
    assert data["summary_status"] == "requires_revision"
    assert data["total_issues"] > 0

    issue_types = {
        issue["issue_type"]
        for issue in data["critical"] + data["major"] + data["minor"]
    }

    assert "weak_phrase" in issue_types
    assert "vacancy_requirements_gap" in issue_types

    assert data["vacancy_relevance"] is not None
    assert "fastapi" in data["vacancy_relevance"]["missing_requirements"]
    assert "postgresql" in data["vacancy_relevance"]["missing_requirements"]
    assert "docker" in data["vacancy_relevance"]["missing_requirements"]

    assert data["technical_info"]["metadata"]["semantic_checks_enabled"] is True
    assert data["technical_info"]["metadata"]["vacancy_text_provided"] is True