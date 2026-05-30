from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from tests.auth_helpers import admin_auth_headers, auth_headers


client = TestClient(app)


def test_user_cannot_read_another_users_report(tmp_path: Path) -> None:
    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, FastAPI, PostgreSQL")
    document.add_paragraph("Опыт работы: Python-разработчик, 2021-2024")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    owner_headers = auth_headers(client, "candidate")
    another_user_headers = auth_headers(client, "candidate")

    with file_path.open("rb") as file:
        create_response = client.post(
            "/api/v1/documents/report",
            headers=owner_headers,
            data={
                "storage_mode": "temporary",
            },
            files={
                "file": (
                    "resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert create_response.status_code == 200

    report_id = create_response.json()["report_id"]

    owner_response = client.get(
        f"/api/v1/documents/reports/{report_id}",
        headers=owner_headers,
    )

    assert owner_response.status_code == 200

    another_user_response = client.get(
        f"/api/v1/documents/reports/{report_id}",
        headers=another_user_headers,
    )

    assert another_user_response.status_code == 404


def test_admin_can_read_any_report(tmp_path: Path) -> None:
    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, FastAPI, PostgreSQL")
    document.add_paragraph("Опыт работы: Python-разработчик, 2021-2024")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    owner_headers = auth_headers(client, "candidate")

    with file_path.open("rb") as file:
        create_response = client.post(
            "/api/v1/documents/report",
            headers=owner_headers,
            data={
                "storage_mode": "temporary",
            },
            files={
                "file": (
                    "resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert create_response.status_code == 200

    report_id = create_response.json()["report_id"]

    admin_response = client.get(
        f"/api/v1/documents/reports/{report_id}",
        headers=admin_auth_headers(client),
    )

    assert admin_response.status_code == 200

def _create_resume_file(tmp_path: Path, filename: str = "resume.docx") -> Path:
    file_path = tmp_path / filename

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, FastAPI, PostgreSQL")
    document.add_paragraph("Опыт работы: Python-разработчик, 2021-2024")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    return file_path


def _create_report(headers: dict[str, str], file_path: Path) -> str:
    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/report",
            headers=headers,
            data={
                "storage_mode": "temporary",
            },
            files={
                "file": (
                    file_path.name,
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    return response.json()["report_id"]


def test_user_can_list_only_own_reports(tmp_path: Path) -> None:
    owner_headers = auth_headers(client, "candidate")
    another_user_headers = auth_headers(client, "candidate")

    owner_file = _create_resume_file(tmp_path, "owner_resume.docx")
    another_file = _create_resume_file(tmp_path, "another_resume.docx")

    owner_report_id = _create_report(owner_headers, owner_file)
    another_report_id = _create_report(another_user_headers, another_file)

    response = client.get(
        "/api/v1/documents/reports",
        headers=owner_headers,
    )

    assert response.status_code == 200

    data = response.json()
    report_ids = {
        item["report_id"]
        for item in data["reports"]
    }

    assert owner_report_id in report_ids
    assert another_report_id not in report_ids
    assert data["total"] == len(data["reports"])


def test_admin_can_list_all_reports(tmp_path: Path) -> None:
    first_user_headers = auth_headers(client, "candidate")
    second_user_headers = auth_headers(client, "hr")

    first_file = _create_resume_file(tmp_path, "first_resume.docx")
    second_file = _create_resume_file(tmp_path, "second_resume.docx")

    first_report_id = _create_report(first_user_headers, first_file)
    second_report_id = _create_report(second_user_headers, second_file)

    response = client.get(
        "/api/v1/documents/reports",
        headers=admin_auth_headers(client),
    )

    assert response.status_code == 200

    data = response.json()
    report_ids = {
        item["report_id"]
        for item in data["reports"]
    }

    assert first_report_id in report_ids
    assert second_report_id in report_ids


def test_list_reports_limit_is_validated() -> None:
    headers = auth_headers(client, "candidate")

    too_small_response = client.get(
        "/api/v1/documents/reports?limit=0",
        headers=headers,
    )

    too_large_response = client.get(
        "/api/v1/documents/reports?limit=101",
        headers=headers,
    )

    assert too_small_response.status_code == 422
    assert too_large_response.status_code == 422