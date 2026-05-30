from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from tests.auth_helpers import auth_headers


client = TestClient(app)


def _create_resume_file(tmp_path: Path) -> Path:
    file_path = tmp_path / "rag_fallback_resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, Git")
    document.add_paragraph("Опыт работы: Backend developer, 2023-2024")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    return file_path


def test_report_generation_falls_back_when_personal_rag_index_missing(
    tmp_path: Path,
) -> None:
    headers = auth_headers(client, "hr")
    file_path = _create_resume_file(tmp_path)

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/report",
            headers=headers,
            data={
                "vacancy_text": "Требования: Python, FastAPI, PostgreSQL, Docker, Git.",
                "storage_mode": "temporary",
            },
            files={
                "file": (
                    file_path.name,
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200

    data = response.json()
    metadata = data["technical_info"]["metadata"]

    assert metadata["rag_backend"] == "per_user_faiss"
    assert metadata["rag_context_used"] is False
    assert metadata["rag_reindex_required"] is True
    assert metadata["rag_index_status"] in {
        "missing",
        "stale",
        "building",
        "failed",
    }
    assert metadata["rag_results_count"] == 0
    assert metadata["rag_reindex_endpoint"] == "/api/v1/rag/reindex"


def test_web_report_generation_falls_back_when_personal_rag_index_missing(
    tmp_path: Path,
) -> None:
    headers = auth_headers(client, "hr")
    file_path = _create_resume_file(tmp_path)

    login_token = headers["Authorization"].replace("Bearer ", "")

    client.cookies.set(
        "access_token",
        login_token,
    )

    try:
        with file_path.open("rb") as file:
            response = client.post(
                "/web/report",
                data={
                    "vacancy_text": "Требования: Python, FastAPI, PostgreSQL, Docker, Git.",
                    "storage_mode": "temporary",
                },
                files={
                    "file": (
                        file_path.name,
                        file,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

        assert response.status_code == 200
        assert "Результат проверки" in response.text

    finally:
        client.cookies.clear()