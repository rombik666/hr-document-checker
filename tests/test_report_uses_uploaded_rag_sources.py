from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.schemas.rag import RagContext
from tests.auth_helpers import auth_headers


client = TestClient(app)


def _create_resume_file(tmp_path: Path) -> Path:
    file_path = tmp_path / "rag_report_resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, Git")
    document.add_paragraph("Опыт работы: Backend developer, 2023-2024")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    return file_path


def test_report_generation_uses_per_user_faiss_rag_for_hr(
    tmp_path: Path,
    monkeypatch,
) -> None:
    calls: list[dict[str, str]] = []

    def fake_search_user_index(
        self,
        owner_user_id: str,
        request,
    ) -> RagContext:
        calls.append(
            {
                "query": request.query,
                "owner_user_id": owner_user_id,
            }
        )

        return RagContext(
            query=request.query,
            results=[],
        )

    monkeypatch.setattr(
        "app.services.rag_index_service.RagIndexService.search_user_index",
        fake_search_user_index,
    )

    headers = auth_headers(client, "hr")
    file_path = _create_resume_file(tmp_path)

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/report",
            headers=headers,
            data={
                "vacancy_text": "Требования: Python, FastAPI, PostgreSQL, Docker, Git.",
                "storage_mode": "temporary",
            },
            files={
                "file": (
                    file_path.name,
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200

    data = response.json()
    metadata = data["technical_info"]["metadata"]

    assert len(calls) == 1
    assert calls[0]["owner_user_id"]
    assert "качество резюме" in calls[0]["query"]
    assert "FastAPI" in calls[0]["query"]
    assert "PostgreSQL" in calls[0]["query"]

    assert metadata["rag_backend"] == "per_user_faiss"
    assert metadata["rag_index_status"] == "ready"
    assert metadata["rag_reindex_required"] is False


def test_report_generation_does_not_use_corporate_rag_for_candidate(
    tmp_path: Path,
    monkeypatch,
) -> None:
    calls: list[dict[str, str]] = []

    def fake_search_user_index(
        self,
        owner_user_id: str,
        request,
    ) -> RagContext:
        calls.append(
            {
                "query": request.query,
                "owner_user_id": owner_user_id,
            }
        )

        return RagContext(
            query=request.query,
            results=[],
        )

    monkeypatch.setattr(
        "app.services.rag_index_service.RagIndexService.search_user_index",
        fake_search_user_index,
    )

    headers = auth_headers(client, "candidate")
    file_path = _create_resume_file(tmp_path)

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/report",
            headers=headers,
            data={
                "vacancy_text": "Требования: Python, FastAPI, PostgreSQL, Docker, Git.",
                "storage_mode": "temporary",
            },
            files={
                "file": (
                    file_path.name,
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200

    data = response.json()
    metadata = data["technical_info"]["metadata"]

    assert calls == []
    assert metadata["rag_backend"] == "per_user_faiss"
    assert metadata["rag_context_used"] is False
    assert metadata["rag_index_status"] == "not_applicable"