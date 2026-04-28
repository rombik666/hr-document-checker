from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


client = TestClient(app)


def test_check_semantic_endpoint_returns_issues(tmp_path: Path) -> None:
    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, Git")
    document.add_paragraph("Опыт работы:")
    document.add_paragraph("Backend developer, 2023-2024")
    document.add_paragraph("Занимался разработкой backend.")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    with file_path.open("rb") as file:
        response = client.post(
            "/api/v1/documents/check-semantic",
            data={
                "vacancy_text": "Требования: Python, FastAPI, PostgreSQL, Docker, Git.",
                "storage_mode": "temporary",
            },
            files={
                "file": (
                    "resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200

    data = response.json()

    assert data["filename"] == "resume.docx"
    assert data["total_issues"] > 0
    assert len(data["check_results"]) == 3