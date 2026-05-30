from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.db.models import IssueORM, ProcessingSessionORM, ReportORM
from app.db.session import SessionLocal
from app.main import app
from tests.auth_helpers import auth_headers


client = TestClient(app)


def _create_resume_file(tmp_path: Path) -> Path:
    file_path = tmp_path / "resume_with_vacancy.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, Git")
    document.add_paragraph("Опыт работы:")
    document.add_paragraph("Backend developer, 2023-2024")
    document.add_paragraph("Занимался разработкой backend.")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    return file_path


def _extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))

    return "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )


def test_resume_with_vacancy_is_saved_to_db_and_exported_to_docx(tmp_path: Path) -> None:
    file_path = _create_resume_file(tmp_path)

    headers = auth_headers(client, "candidate")

    with file_path.open("rb") as file:
        create_response = client.post(
            "/api/v1/documents/report",
            headers=headers,
            data={
                "vacancy_text": "Требования: Python, FastAPI, PostgreSQL, Docker, Git.",
                "storage_mode": "temporary",
            },
            files={
                "file": (
                    file_path.name,
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert create_response.status_code == 200

    report_data = create_response.json()

    report_id = report_data["report_id"]
    document_id = report_data["document_id"]

    assert report_data["vacancy_relevance"] is not None
    assert "fastapi" in report_data["vacancy_relevance"]["missing_requirements"]
    assert "postgresql" in report_data["vacancy_relevance"]["missing_requirements"]
    assert "docker" in report_data["vacancy_relevance"]["missing_requirements"]

    db = SessionLocal()

    try:
        stored_report = db.get(ReportORM, report_id)

        assert stored_report is not None
        assert stored_report.document_id == document_id
        assert stored_report.report_json["vacancy_relevance"] is not None
        assert "fastapi" in stored_report.report_json["vacancy_relevance"]["missing_requirements"]
        assert "postgresql" in stored_report.report_json["vacancy_relevance"]["missing_requirements"]
        assert "docker" in stored_report.report_json["vacancy_relevance"]["missing_requirements"]

        processing_session = (
            db.query(ProcessingSessionORM)
            .filter(ProcessingSessionORM.document_id == document_id)
            .one()
        )

        assert processing_session.session_metadata["vacancy_relevance_present"] is True
        assert processing_session.session_metadata["report_id"] == report_id

        vacancy_issue = (
            db.query(IssueORM)
            .filter(IssueORM.document_id == document_id)
            .filter(IssueORM.report_id == report_id)
            .filter(IssueORM.issue_type == "vacancy_requirements_gap")
            .one_or_none()
        )

        assert vacancy_issue is not None
        assert vacancy_issue.severity == "Critical"
        assert vacancy_issue.evidence_fragment is not None
        assert "fastapi" in vacancy_issue.evidence_fragment
        assert "postgresql" in vacancy_issue.evidence_fragment
        assert "docker" in vacancy_issue.evidence_fragment

    finally:
        db.close()

    export_response = client.get(
        f"/api/v1/documents/reports/{report_id}/export/docx",
        headers=headers,
    )

    assert export_response.status_code == 200
    assert (
        export_response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    exported_text = _extract_docx_text(export_response.content).lower()

    assert "релевантность вакансии" in exported_text
    assert "покрытие требований" in exported_text
    assert "незакрытые требования" in exported_text
    assert "fastapi" in exported_text
    assert "postgresql" in exported_text
    assert "docker" in exported_text