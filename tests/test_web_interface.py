from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from PIL import Image

from app.main import app
from tests.auth_helpers import auth_headers

from sqlalchemy import delete

from app.db.models import RagIndexORM, RagSourceORM
from app.db.session import SessionLocal


client = TestClient(app)


def test_web_index_returns_landing_page_for_anonymous_user() -> None:
    response = TestClient(app).get("/web/")

    assert response.status_code == 200
    assert "HR Document Checker" in response.text
    assert "Проверяйте HR-документы" in response.text
    assert "AI-агенты" in response.text
    assert "RAG" in response.text
    assert "FAISS" in response.text
    assert 'href="/web/register"' in response.text
    assert 'href="/web/login"' in response.text
    assert 'action="/web/login"' not in response.text

def test_web_index_redirects_authenticated_user_to_dashboard() -> None:
    response = client.get(
        "/web/",
        headers=auth_headers(client, "candidate"),
        follow_redirects=False,
    )

    assert response.status_code == 303
    assert response.headers["location"] == "/web/dashboard"


def test_web_dashboard_returns_candidate_dashboard_for_authenticated_user() -> None:
    response = client.get(
        "/web/dashboard",
        headers=auth_headers(client, "candidate"),
    )

    assert response.status_code == 200
    assert "Кабинет кандидата" in response.text
    assert 'href="#check-document"' in response.text
    assert 'action="/web/report"' in response.text
    assert 'name="file"' in response.text
    assert 'name="vacancy_text"' in response.text
    assert 'name="storage_mode"' in response.text
    assert 'href="/web/reports"' in response.text
    assert 'class="user-dropdown-logout"' in response.text
    assert 'src="/static/icons/log-out.svg"' in response.text
    assert "/web/rag/sources" not in response.text
    assert "/web/admin" not in response.text


def test_profile_edit_updates_account_and_avatar() -> None:
    headers = auth_headers(client, "candidate")
    image_buffer = BytesIO()
    Image.new("RGB", (2, 2), color="#2563eb").save(image_buffer, format="PNG")
    image_bytes = image_buffer.getvalue()

    upload_response = client.post(
        "/web/profile/edit",
        headers=headers,
        data={
            "full_name": "Обновлённый Пользователь",
            "email": f"updated-{id(headers)}@example.com",
        },
        files={"avatar": ("avatar.png", image_bytes, "image/png")},
        follow_redirects=False,
    )

    assert upload_response.status_code == 303
    assert upload_response.headers["location"] == "/web/profile?updated=true"

    avatar_response = client.get("/web/profile/avatar", headers=headers)

    assert avatar_response.status_code == 200
    assert avatar_response.headers["content-type"] == "image/png"
    assert avatar_response.content == image_bytes

    profile_response = client.get("/web/profile", headers=headers)

    assert 'src="/web/profile/avatar"' in profile_response.text
    assert "Обновлённый Пользователь" in profile_response.text
    assert 'href="/web/profile/edit"' in profile_response.text
    assert 'class="profile-logout-link"' in profile_response.text


def test_profile_rejects_non_image_avatar() -> None:
    response = client.post(
        "/web/profile/edit",
        headers=auth_headers(client, "candidate"),
        data={
            "full_name": "Candidate Test User",
            "email": f"avatar-error-{id(client)}@example.com",
        },
        files={"avatar": ("avatar.png", b"not-an-image", "image/png")},
    )

    assert response.status_code == 400
    assert "Поддерживаются только изображения PNG, JPEG и WebP." in response.text


def test_anonymous_contact_form_redirects_to_registration() -> None:
    anonymous_client = TestClient(app)
    page_response = anonymous_client.get("/web/")

    assert "Зарегистрируйтесь, чтобы написать нам" in page_response.text
    assert 'data-contact-form' not in page_response.text

    response = anonymous_client.post(
        "/web/contact",
        data={"topic": "Вопрос", "message": "Сообщение"},
    )

    assert response.status_code == 401
    assert response.json()["redirect_url"] == "/web/register"


def test_web_report_page_returns_html_report(tmp_path: Path) -> None:
    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, Git")
    document.add_paragraph("Опыт работы:")
    document.add_paragraph("Backend developer, 2023-2024")
    document.add_paragraph("Занимался разработкой backend.")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    with file_path.open("rb") as file:
        response = client.post(
            "/web/report",
            headers=auth_headers(client, "candidate"),
            data={
                "vacancy_text": "Требования: Python, FastAPI, PostgreSQL, Docker, Git.",
                "storage_mode": "no_store",
            },
            files={
                "file": (
                    "resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    assert "Результат проверки" in response.text
    assert "Всего замечаний" in response.text
    assert "Новая проверка" in response.text


def test_web_report_rejects_unsupported_file_format(tmp_path: Path) -> None:
    file_path = tmp_path / "resume.txt"
    file_path.write_text("Simple text", encoding="utf-8")

    with file_path.open("rb") as file:
        response = client.post(
            "/web/report",
            headers=auth_headers(client, "candidate"),
            data={
                "storage_mode": "temporary",
            },
            files={
                "file": (
                    "resume.txt",
                    file,
                    "text/plain",
                )
            },
        )

    assert response.status_code == 400
    assert "Поддерживаются только файлы .docx и .pdf" in response.text

def test_web_reports_history_requires_authentication() -> None:
    anonymous_client = TestClient(app)

    response = anonymous_client.get(
        "/web/reports",
        follow_redirects=False,
    )

    assert response.status_code == 303
    assert response.headers["location"] == "/web/login"


def test_web_reports_history_returns_authenticated_user_reports() -> None:
    response = client.get(
        "/web/reports",
        headers=auth_headers(client, "candidate"),
    )

    assert response.status_code == 200
    assert "История проверок" in response.text


def test_web_saved_report_page_returns_html_report(tmp_path: Path) -> None:
    file_path = tmp_path / "saved_resume.docx"

    document = Document()
    document.add_paragraph("Контакты:")
    document.add_paragraph("Email: ivan@example.com")
    document.add_paragraph("Телефон: +7 999 123-45-67")
    document.add_paragraph("Навыки: Python, Git")
    document.add_paragraph("Опыт работы: Backend developer, 2023-2024")
    document.add_paragraph("Образование: Южный федеральный университет")
    document.save(file_path)

    headers = auth_headers(client, "candidate")

    with file_path.open("rb") as file:
        create_response = client.post(
            "/web/report",
            headers=headers,
            data={
                "vacancy_text": "Требования: Python, FastAPI, PostgreSQL, Docker, Git.",
                "storage_mode": "temporary",
            },
            files={
                "file": (
                    "saved_resume.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert create_response.status_code == 200

    # report_id есть в ссылке DOCX-экспорта сохранённого отчёта.
    marker = "/api/v1/documents/reports/"
    assert marker in create_response.text

    report_id = (
        create_response.text
        .split(marker, 1)[1]
        .split("/export/docx", 1)[0]
    )

    history_response = client.get(
        "/web/reports",
        headers=headers,
    )

    assert history_response.status_code == 200
    assert f'href="/web/reports/{report_id}"' in history_response.text
    assert f'href="/api/v1/documents/reports/{report_id}/export/docx"' in history_response.text
    assert "/web/reports/" in history_response.text
    assert "/api/v1/documents/reports//export/docx" not in history_response.text

    assert report_id

    saved_response = client.get(
        f"/web/reports/{report_id}",
        headers=headers,
    )

    assert saved_response.status_code == 200
    assert "Результат проверки" in saved_response.text
    assert "Скачать DOCX" in saved_response.text
    assert "Релевантность вакансии" in saved_response.text

def _cleanup_web_rag_sources() -> None:
    db = SessionLocal()

    try:
        sources = (
            db.query(RagSourceORM)
            .filter(RagSourceORM.filename.like("%web_rag_ui_test%"))
            .all()
        )

        owner_user_ids = {
            source.owner_user_id
            for source in sources
            if source.owner_user_id is not None
        }

        db.execute(
            delete(RagSourceORM).where(
                RagSourceORM.filename.like("%web_rag_ui_test%")
            )
        )

        for owner_user_id in owner_user_ids:
            db.execute(
                delete(RagIndexORM).where(
                    RagIndexORM.owner_user_id == owner_user_id,
                )
            )

        db.commit()

    finally:
        db.close()


def test_web_rag_sources_requires_authentication() -> None:
    anonymous_client = TestClient(app)

    response = anonymous_client.get(
        "/web/rag/sources",
        follow_redirects=False,
    )

    assert response.status_code == 303
    assert response.headers["location"] == "/web/login"


def test_web_rag_sources_forbids_candidate() -> None:
    response = client.get(
        "/web/rag/sources",
        headers=auth_headers(client, "candidate"),
    )

    assert response.status_code == 403
    assert "RAG-источники доступны только HR-специалистам" in response.text


def test_web_rag_sources_page_returns_hr_ui() -> None:
    response = client.get(
        "/web/rag/sources",
        headers=auth_headers(client, "hr"),
    )

    assert response.status_code == 200
    assert "RAG-источники" in response.text
    assert "Загрузить источник" in response.text
    assert "Загруженные источники" in response.text
    assert "Статус FAISS-индекса" in response.text
    assert "Переиндексировать" in response.text


def test_web_hr_can_upload_rag_source(tmp_path: Path) -> None:
    _cleanup_web_rag_sources()

    headers = auth_headers(client, "hr")
    file_path = tmp_path / "web_rag_ui_test_requirements.txt"

    file_path.write_text(
        "Требования компании: Python, FastAPI, PostgreSQL, Docker.",
        encoding="utf-8",
    )

    try:
        with file_path.open("rb") as file:
            response = client.post(
                "/web/rag/sources/upload",
                headers=headers,
                data={
                    "title": "Web RAG UI source",
                    "source_type": "requirements",
                },
                files={
                    "file": (
                        file_path.name,
                        file,
                        "text/plain",
                    )
                },
            )

        assert response.status_code == 200
        assert "RAG-источник успешно загружен" in response.text
        assert "Web RAG UI source" in response.text
        assert "web_rag_ui_test_requirements.txt" in response.text

    finally:
        _cleanup_web_rag_sources()


def test_web_hr_can_deactivate_rag_source(tmp_path: Path) -> None:
    _cleanup_web_rag_sources()

    headers = auth_headers(client, "hr")
    file_path = tmp_path / "web_rag_ui_test_delete.txt"

    file_path.write_text(
        "Источник для отключения через Web UI.",
        encoding="utf-8",
    )

    try:
        with file_path.open("rb") as file:
            upload_response = client.post(
                "/web/rag/sources/upload",
                headers=headers,
                data={
                    "title": "Web RAG UI delete source",
                    "source_type": "other",
                },
                files={
                    "file": (
                        file_path.name,
                        file,
                        "text/plain",
                    )
                },
            )

        assert upload_response.status_code == 200

        db = SessionLocal()

        try:
            source = (
                db.query(RagSourceORM)
                .filter(RagSourceORM.filename == "web_rag_ui_test_delete.txt")
                .one()
            )

            source_id = source.id

        finally:
            db.close()

        delete_response = client.post(
            f"/web/rag/sources/{source_id}/delete",
            headers=headers,
        )

        assert delete_response.status_code == 200
        assert "RAG-источник отключён" in delete_response.text
        assert "inactive" in delete_response.text

    finally:
        _cleanup_web_rag_sources()


def test_web_rag_source_upload_rejects_unsupported_format(tmp_path: Path) -> None:
    headers = auth_headers(client, "hr")
    file_path = tmp_path / "web_rag_ui_test_bad.exe"

    file_path.write_text("bad file", encoding="utf-8")

    with file_path.open("rb") as file:
        response = client.post(
            "/web/rag/sources/upload",
            headers=headers,
            data={
                "title": "Bad source",
                "source_type": "other",
            },
            files={
                "file": (
                    file_path.name,
                    file,
                    "application/octet-stream",
                )
            },
        )

    assert response.status_code == 400
    assert "Поддерживаются только RAG-источники" in response.text

def test_web_hr_can_activate_deactivated_rag_source(tmp_path: Path) -> None:
    _cleanup_web_rag_sources()

    headers = auth_headers(client, "hr")
    file_path = tmp_path / "web_rag_ui_test_activate.txt"

    file_path.write_text(
        "Источник для включения через Web UI.",
        encoding="utf-8",
    )

    try:
        with file_path.open("rb") as file:
            upload_response = client.post(
                "/web/rag/sources/upload",
                headers=headers,
                data={
                    "title": "Web RAG UI activate source",
                    "source_type": "other",
                },
                files={
                    "file": (
                        file_path.name,
                        file,
                        "text/plain",
                    )
                },
            )

        assert upload_response.status_code == 200

        db = SessionLocal()

        try:
            source = (
                db.query(RagSourceORM)
                .filter(RagSourceORM.filename == "web_rag_ui_test_activate.txt")
                .one()
            )

            source_id = source.id

        finally:
            db.close()

        deactivate_response = client.post(
            f"/web/rag/sources/{source_id}/delete",
            headers=headers,
        )

        assert deactivate_response.status_code == 200
        assert "RAG-источник отключён" in deactivate_response.text

        activate_response = client.post(
            f"/web/rag/sources/{source_id}/activate",
            headers=headers,
        )

        assert activate_response.status_code == 200
        assert "RAG-источник включён" in activate_response.text
        assert "active" in activate_response.text

    finally:
        _cleanup_web_rag_sources()


def test_web_hr_can_permanently_delete_rag_source(tmp_path: Path) -> None:
    _cleanup_web_rag_sources()

    headers = auth_headers(client, "hr")
    file_path = tmp_path / "web_rag_ui_test_permanent_delete.txt"

    file_path.write_text(
        "Источник для полного удаления через Web UI.",
        encoding="utf-8",
    )

    try:
        with file_path.open("rb") as file:
            upload_response = client.post(
                "/web/rag/sources/upload",
                headers=headers,
                data={
                    "title": "Web RAG UI permanent delete source",
                    "source_type": "other",
                },
                files={
                    "file": (
                        file_path.name,
                        file,
                        "text/plain",
                    )
                },
            )

        assert upload_response.status_code == 200

        db = SessionLocal()

        try:
            source = (
                db.query(RagSourceORM)
                .filter(RagSourceORM.filename == "web_rag_ui_test_permanent_delete.txt")
                .one()
            )

            source_id = source.id

        finally:
            db.close()

        delete_response = client.post(
            f"/web/rag/sources/{source_id}/permanent-delete",
            headers=headers,
        )

        assert delete_response.status_code == 200
        assert "RAG-источник полностью удалён" in delete_response.text
        assert "Web RAG UI permanent delete source" not in delete_response.text

        db = SessionLocal()

        try:
            deleted_source = db.get(RagSourceORM, source_id)
            assert deleted_source is None

        finally:
            db.close()

    finally:
        _cleanup_web_rag_sources()

def test_web_hr_can_reindex_rag_sources(tmp_path: Path) -> None:
    _cleanup_web_rag_sources()

    headers = auth_headers(client, "hr")
    file_path = tmp_path / "web_rag_ui_test_reindex.txt"

    file_path.write_text(
        "Источник для переиндексации через Web UI: Python, FastAPI, PostgreSQL.",
        encoding="utf-8",
    )

    try:
        with file_path.open("rb") as file:
            upload_response = client.post(
                "/web/rag/sources/upload",
                headers=headers,
                data={
                    "title": "Web RAG UI reindex source",
                    "source_type": "requirements",
                },
                files={
                    "file": (
                        file_path.name,
                        file,
                        "text/plain",
                    )
                },
            )

        assert upload_response.status_code == 200
        assert "stale" in upload_response.text
        assert "Переиндексировать" in upload_response.text

        reindex_response = client.post(
            "/web/rag/reindex",
            headers=headers,
        )

        assert reindex_response.status_code == 200
        assert "RAG-индекс переиндексирован" in reindex_response.text
        assert "ready" in reindex_response.text
        assert "Web RAG UI reindex source" in reindex_response.text

    finally:
        _cleanup_web_rag_sources()
